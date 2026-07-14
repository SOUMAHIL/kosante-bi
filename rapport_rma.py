# =============================================================
# KoSanté BI — rapport_rma.py
# =============================================================
# Génération automatique du Rapport Mensuel des Activités (RMA)
# JHPIEGO RISE — Suivi PVVIH
#
# USAGE :
#   python rapport_rma.py --mois 2026-06
# =============================================================

import duckdb
import pandas as pd
from pathlib import Path
from datetime import date, datetime
from dateutil.relativedelta import relativedelta
import argparse
import logging

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(message)s",
    datefmt="%H:%M:%S"
)
log = logging.getLogger(__name__)

# =============================================================
# CONFIG
# =============================================================
CONFIG_ETAB = {
    "nom":      "ONG KO'KHOUA",
    "district": "Treichville-Marcory",
    "region":   "Abidjan 2",
}

# Tranches d'âge RMA — 15 tranches
TRANCHES_RMA = [
    ("lt1",   "< 1 an",       "Age < 1"),
    ("t14",   "1-4 ans",      "Age BETWEEN 1 AND 4"),
    ("t59",   "5-9 ans",      "Age BETWEEN 5 AND 9"),
    ("t1014", "10-14 ans",    "Age BETWEEN 10 AND 14"),
    ("t1519", "15-19 ans",    "Age BETWEEN 15 AND 19"),
    ("t2024", "20-24 ans",    "Age BETWEEN 20 AND 24"),
    ("t2529", "25-29 ans",    "Age BETWEEN 25 AND 29"),
    ("t3034", "30-34 ans",    "Age BETWEEN 30 AND 34"),
    ("t3539", "35-39 ans",    "Age BETWEEN 35 AND 39"),
    ("t4044", "40-44 ans",    "Age BETWEEN 40 AND 44"),
    ("t4549", "45-49 ans",    "Age BETWEEN 45 AND 49"),
    ("t5054", "50-54 ans",    "Age BETWEEN 50 AND 54"),
    ("t5559", "55-59 ans",    "Age BETWEEN 55 AND 59"),
    ("t6064", "60-64 ans",    "Age BETWEEN 60 AND 64"),
    ("t65p",  "65 ans et +",  "Age >= 65"),
]

# Régimes ARV 1ère ligne
REGIMES_1ERE_LIGNE = [
    "ABC 3TC DTG", "ABC 3TC EFV", "ABC 3TC LPV/r", "ABC 3TC NVP",
    "ATZ/r 3TC DTG", "AZT 3TC ABC", "AZT 3TC ATV/r", "AZT 3TC DTG",
    "AZT 3TC EFV", "AZT 3TC LPV/r", "AZT 3TC NVP", "AZT 3TC TDF",
    "TDF 3TC ATV/r", "TDF 3TC DTG", "TDF 3TC DTG DTG", "TDF 3TC EFV",
    "TDF 3TC LPV/r", "TDF 3TC NVP", "TDF FTC EFV", "AUTRE"
]

# Régimes ARV 2ème ligne
REGIMES_2EME_LIGNE = [
    "ABC 3TC DTG", "ABC 3TC EFV", "ABC 3TC LPV/r", "ABC 3TC NVP",
    "ATZ/r 3TC DTG", "AZT 3TC ABC", "AZT 3TC ATV/r", "AZT 3TC DTG",
    "AZT 3TC EFV", "AZT 3TC LPV/r", "AZT 3TC NVP", "AZT 3TC TDF",
    "TDF 3TC ATV/r", "TDF 3TC DTG", "TDF 3TC DTG DTG", "TDF 3TC EFV",
    "TDF 3TC LPV/r", "TDF 3TC NVP", "TDF FTC EFV", "AUTRE"
]


# =============================================================
# CONNEXION DUCKDB
# =============================================================
def get_conn():
    db_path = Path(__file__).parent / "data" / "kosante.duckdb"
    return duckdb.connect(str(db_path), read_only=True)


def get_periode(mois_str):
    d = datetime.strptime(mois_str, "%Y-%m")
    debut = date(d.year, d.month, 1)
    fin = debut + relativedelta(months=1) - relativedelta(days=1)
    label = d.strftime("%B %Y").capitalize()
    return debut, fin, label


# =============================================================
# SQL HELPERS
# =============================================================
NORMALISE = """REGEXP_REPLACE(TRIM(CAST({col} AS VARCHAR)), '^0+', '')"""


def sql_rma_sexe(champ_age, champ_sexe, filtre_extra="", prefix=""):
    """Génère colonnes SQL pour les 15 tranches RMA × M/F"""
    cols = []
    p = f"{prefix}_" if prefix else ""
    for code, _, cond in TRANCHES_RMA:
        full = f"({cond})"
        if filtre_extra:
            full = f"({cond}) AND ({filtre_extra})"
        cols.append(
            f"COUNT(*) FILTER (WHERE {full} AND {champ_sexe} = 1) AS {p}m_{code}"
        )
        cols.append(
            f"COUNT(*) FILTER (WHERE {full} AND {champ_sexe} = 2) AS {p}f_{code}"
        )
    extra_cond = f"AND ({filtre_extra})" if filtre_extra else ""
    cols.append(f"COUNT(*) FILTER (WHERE {champ_sexe} = 1 {extra_cond}) AS {p}m_total")
    cols.append(f"COUNT(*) FILTER (WHERE {champ_sexe} = 2 {extra_cond}) AS {p}f_total")
    cols.append(f"COUNT(*) {('FILTER (WHERE ' + filtre_extra + ')') if filtre_extra else ''} AS {p}total")
    return cols


def file_active_fin_mois(fin, filtre_extra=""):
    """SQL pour file active figée à la fin du mois"""
    extra = f"AND ({filtre_extra})" if filtre_extra else ""
    return f"""
        SELECT p.Age, p.Sexe, r.JOURS, r.REGIME AS DernierRegime,
               sc.StatutCV
        FROM TblRegime r
        INNER JOIN TblDossPatient p
            ON {NORMALISE.format(col='r.NumPatient')}
             = {NORMALISE.format(col='p.NumNational')}
        LEFT JOIN StatutCV_Patient sc ON p.NumInc = sc.NumInc
        WHERE r.Derniere_Dispensation = TRUE
          AND r.REGIME IS NOT NULL
          AND TRY_CAST(r.DatePDV AS DATE) >= DATE '{fin}'
          AND (p.DecesDate IS NULL OR p.DecesDate = '')
          AND (p.TransfDate IS NULL OR p.TransfDate = '')
          AND p.NumInc > 0
          {extra}
    """


def val(row, col, default=0):
    try:
        v = row[col]
        if pd.isna(v):
            return default
        return int(v)
    except (KeyError, TypeError):
        return default


# =============================================================
# REQUÊTES
# =============================================================

def q_screening(conn, debut, fin):
    """Screening VIH — tableau consultants"""

    filtre_screen = """
        (NumPartenaire IS NULL OR TRIM(CAST(NumPartenaire AS VARCHAR)) = '')
        AND (NumParent IS NULL OR TRIM(CAST(NumParent AS VARCHAR)) = '')
    """

    # Sous-requête avec âge pré-calculé
    def base_cdv(filtre_extra=""):
        extra = f"AND ({filtre_extra})" if filtre_extra else ""
        return f"""
            SELECT
                DATEDIFF('year',
                    TRY_CAST(DateNaiss AS DATE),
                    TRY_CAST(DateVisite AS DATE)
                ) AS Age,
                Sexe
            FROM TblRegistreCDV
            WHERE TRY_CAST(DateVisite AS DATE)
                  BETWEEN DATE '{debut}' AND DATE '{fin}'
              AND DateNaiss IS NOT NULL
              {extra}
        """

    tous = conn.execute(f"""
        SELECT {', '.join(sql_rma_sexe('Age', 'Sexe'))}
        FROM ({base_cdv()}) t
    """).df().iloc[0]

    screens = conn.execute(f"""
        SELECT {', '.join(sql_rma_sexe('Age', 'Sexe'))}
        FROM ({base_cdv(filtre_screen)}) t
    """).df().iloc[0]

    positifs = conn.execute(f"""
        SELECT {', '.join(sql_rma_sexe('Age', 'Sexe'))}
        FROM (
            SELECT
                DATEDIFF('year',
                    TRY_CAST(DateNaiss AS DATE),
                    TRY_CAST(DateVisite AS DATE)
                ) AS Age, Sexe
            FROM TblRegistreCDV
            WHERE TRY_CAST(DateVisite AS DATE)
                  BETWEEN DATE '{debut}' AND DATE '{fin}'
              AND DateNaiss IS NOT NULL
              AND ({filtre_screen})
              AND Resultat_simple = 'Positif'
        ) t
    """).df().iloc[0]

    return {"tous": tous, "screens": screens, "positifs": positifs}


def q_conseil_depistage(conn, debut, fin):
    """I- Conseil et Dépistage"""

    filtre_screen = """
        (NumPartenaire IS NULL OR TRIM(CAST(NumPartenaire AS VARCHAR)) = '')
        AND (NumParent IS NULL OR TRIM(CAST(NumParent AS VARCHAR)) = '')
    """

    def base_cdv(filtre_extra=""):
        extra = f"AND ({filtre_extra})" if filtre_extra else ""
        return f"""
            SELECT
                DATEDIFF('year',
                    TRY_CAST(DateNaiss AS DATE),
                    TRY_CAST(DateVisite AS DATE)
                ) AS Age,
                Sexe
            FROM TblRegistreCDV
            WHERE TRY_CAST(DateVisite AS DATE)
                  BETWEEN DATE '{debut}' AND DATE '{fin}'
              AND DateNaiss IS NOT NULL
              {extra}
        """

    tous = conn.execute(f"""
        SELECT {', '.join(sql_rma_sexe('Age', 'Sexe'))}
        FROM ({base_cdv(filtre_screen)}) t
    """).df().iloc[0]

    positifs = conn.execute(f"""
        SELECT {', '.join(sql_rma_sexe('Age', 'Sexe'))}
        FROM (
            SELECT
                DATEDIFF('year',
                    TRY_CAST(DateNaiss AS DATE),
                    TRY_CAST(DateVisite AS DATE)
                ) AS Age, Sexe
            FROM TblRegistreCDV
            WHERE TRY_CAST(DateVisite AS DATE)
                  BETWEEN DATE '{debut}' AND DATE '{fin}'
              AND DateNaiss IS NOT NULL
              AND ({filtre_screen})
              AND Resultat_simple = 'Positif'
        ) t
    """).df().iloc[0]

    negatifs = conn.execute(f"""
        SELECT {', '.join(sql_rma_sexe('Age', 'Sexe'))}
        FROM (
            SELECT
                DATEDIFF('year',
                    TRY_CAST(DateNaiss AS DATE),
                    TRY_CAST(DateVisite AS DATE)
                ) AS Age, Sexe
            FROM TblRegistreCDV
            WHERE TRY_CAST(DateVisite AS DATE)
                  BETWEEN DATE '{debut}' AND DATE '{fin}'
              AND DateNaiss IS NOT NULL
              AND ({filtre_screen})
              AND Resultat_simple = 'Négatif'
        ) t
    """).df().iloc[0]

    return {"tous": tous, "positifs": positifs, "negatifs": negatifs}


def q_file_active_dotation(conn, fin):
    """312 — File active par dotation"""

    base = file_active_fin_mois(fin)

    total = conn.execute(f"""
        SELECT {', '.join(sql_rma_sexe('fa.Age', 'fa.Sexe'))}
        FROM ({base}) fa
    """).df().iloc[0]

    dot_lt90 = conn.execute(f"""
        SELECT {', '.join(sql_rma_sexe('fa.Age', 'fa.Sexe', 'fa.JOURS < 90'))}
        FROM ({base}) fa
    """).df().iloc[0]

    dot_90_179 = conn.execute(f"""
        SELECT {', '.join(sql_rma_sexe('fa.Age', 'fa.Sexe',
               'fa.JOURS BETWEEN 90 AND 179'))}
        FROM ({base}) fa
    """).df().iloc[0]

    dot_gte180 = conn.execute(f"""
        SELECT {', '.join(sql_rma_sexe('fa.Age', 'fa.Sexe', 'fa.JOURS >= 180'))}
        FROM ({base}) fa
    """).df().iloc[0]

    # FE (Femmes enceintes) et FA (Femmes allaitantes) → vide dans notre BD
    return {
        "total":      total,
        "dot_lt90":   dot_lt90,
        "dot_90_179": dot_90_179,
        "dot_gte180": dot_gte180,
    }


def q_stables(conn, fin):
    """318 — Patients stables par modèle"""

    base = file_active_fin_mois(fin, "sc.StatutCV = 'Stable'")

    total = conn.execute(f"""
        SELECT {', '.join(sql_rma_sexe('fa.Age', 'fa.Sexe'))}
        FROM ({base}) fa
    """).df().iloc[0]

    # 318.a : stables JOURS <= 90
    a = conn.execute(f"""
        SELECT {', '.join(sql_rma_sexe('fa.Age', 'fa.Sexe', 'fa.JOURS <= 90'))}
        FROM ({base}) fa
    """).df().iloc[0]

    # 318.b : stables JOURS >= 180
    b = conn.execute(f"""
        SELECT {', '.join(sql_rma_sexe('fa.Age', 'fa.Sexe', 'fa.JOURS >= 180'))}
        FROM ({base}) fa
    """).df().iloc[0]

    # 318.f : stables JOURS BETWEEN 91 AND 179
    f = conn.execute(f"""
        SELECT {', '.join(sql_rma_sexe('fa.Age', 'fa.Sexe',
               'fa.JOURS BETWEEN 91 AND 179'))}
        FROM ({base}) fa
    """).df().iloc[0]

    # 321 Non stables
    base_ns = file_active_fin_mois(fin, "sc.StatutCV = 'Non Stable'")
    non_stables = conn.execute(f"""
        SELECT {', '.join(sql_rma_sexe('fa.Age', 'fa.Sexe'))}
        FROM ({base_ns}) fa
    """).df().iloc[0]

    # 324 Non évalués
    base_ne = file_active_fin_mois(fin, "sc.StatutCV = 'NE'")
    non_evalues = conn.execute(f"""
        SELECT {', '.join(sql_rma_sexe('fa.Age', 'fa.Sexe'))}
        FROM ({base_ne}) fa
    """).df().iloc[0]

    return {
        "total":       total,
        "a_3mois":     a,
        "b_6mois":     b,
        "f_autres":    f,
        "non_stables": non_stables,
        "non_evalues": non_evalues,
    }


def q_regimes(conn, debut, fin):
    """4 — Régimes ARV"""

    # Nouveaux sous ARV ce mois par régime
    nouveaux = conn.execute(f"""
        SELECT
            r.REGIME AS regime,
            SUM(CASE WHEN p.Sexe = 1 AND p.Age >= 15 THEN 1 ELSE 0 END) AS adulte_h,
            SUM(CASE WHEN p.Sexe = 2 AND p.Age >= 15 THEN 1 ELSE 0 END) AS adulte_f,
            SUM(CASE WHEN p.Sexe = 1 AND p.Age < 15  THEN 1 ELSE 0 END) AS enfant_h,
            SUM(CASE WHEN p.Sexe = 2 AND p.Age < 15  THEN 1 ELSE 0 END) AS enfant_f
        FROM TblMiseEnRoute m
        INNER JOIN TblDossPatient p ON m.Patient = p.NumInc
        INNER JOIN TblRegime r
            ON {NORMALISE.format(col='r.NumPatient')}
             = {NORMALISE.format(col='p.NumNational')}
           AND r.Derniere_Dispensation = TRUE
        WHERE TRY_CAST(m.DateMiseTARV AS DATE) BETWEEN DATE '{debut}' AND DATE '{fin}'
        GROUP BY r.REGIME
    """).df()

    # File active par régime
    file_active_sql = file_active_fin_mois(fin)
    fa_regime = conn.execute(f"""
        SELECT
            fa.DernierRegime AS regime,
            SUM(CASE WHEN fa.Sexe = 1 AND fa.Age >= 15 THEN 1 ELSE 0 END) AS adulte_h,
            SUM(CASE WHEN fa.Sexe = 2 AND fa.Age >= 15 THEN 1 ELSE 0 END) AS adulte_f,
            SUM(CASE WHEN fa.Sexe = 1 AND fa.Age < 15  THEN 1 ELSE 0 END) AS enfant_h,
            SUM(CASE WHEN fa.Sexe = 2 AND fa.Age < 15  THEN 1 ELSE 0 END) AS enfant_f
        FROM ({file_active_sql}) fa
        GROUP BY fa.DernierRegime
    """).df()

    return {"nouveaux": nouveaux, "file_active": fa_regime}


def q_charge_virale(conn, debut, fin):
    """11a — Charge virale 12 derniers mois + mois en cours"""

    debut_12m = (date.fromisoformat(str(fin)) - relativedelta(months=12)).strftime("%Y-%m-%d")

    # 1006 : CV réalisée 12 derniers mois (résultat disponible)
    cv_1006 = conn.execute(f"""
        SELECT {', '.join(sql_rma_sexe('p.Age', 'p.Sexe'))}
        FROM TblChargesVirales cv
        INNER JOIN TblDossPatient p ON cv.Patient = p.NumInc
        WHERE TRY_CAST(cv.DatePrelev AS DATE)
              BETWEEN DATE '{debut_12m}' AND DATE '{fin}'
          AND cv.CVcopies IS NOT NULL
    """).df().iloc[0]

    # 1009 : CV ≤ 1000 copies/ml
    cv_1009 = conn.execute(f"""
        SELECT {', '.join(sql_rma_sexe('p.Age', 'p.Sexe'))}
        FROM TblChargesVirales cv
        INNER JOIN TblDossPatient p ON cv.Patient = p.NumInc
        WHERE TRY_CAST(cv.DatePrelev AS DATE)
              BETWEEN DATE '{debut_12m}' AND DATE '{fin}'
          AND cv.CVcopies <= 1000
    """).df().iloc[0]

    # 1012 : CV > 1000 copies/ml
    cv_1012 = conn.execute(f"""
        SELECT {', '.join(sql_rma_sexe('p.Age', 'p.Sexe'))}
        FROM TblChargesVirales cv
        INNER JOIN TblDossPatient p ON cv.Patient = p.NumInc
        WHERE TRY_CAST(cv.DatePrelev AS DATE)
              BETWEEN DATE '{debut_12m}' AND DATE '{fin}'
          AND cv.CVcopies > 1000
    """).df().iloc[0]

    # 1022 : CV prélevée CE MOIS
    cv_1022 = conn.execute(f"""
        SELECT {', '.join(sql_rma_sexe('p.Age', 'p.Sexe'))}
        FROM TblChargesVirales cv
        INNER JOIN TblDossPatient p ON cv.Patient = p.NumInc
        WHERE TRY_CAST(cv.DatePrelev AS DATE)
              BETWEEN DATE '{debut}' AND DATE '{fin}'
          AND cv.CVcopies IS NOT NULL
    """).df().iloc[0]

    return {
        "cv_1006": cv_1006,
        "cv_1009": cv_1009,
        "cv_1012": cv_1012,
        "cv_1022": cv_1022,
    }


# =============================================================
# HELPERS DOCX
# =============================================================
def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def cell_text(cell, text, bold=False, size=7,
              align=WD_ALIGN_PARAGRAPH.CENTER, color=None, bg=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(str(text) if text is not None else "")
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    if bg:
        set_cell_bg(cell, bg)


def add_titre(doc, text, level=1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12 if level == 1 else 10)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)


def build_header_rma(t, titre_col0=""):
    """Construit l'en-tête standard RMA : libellé | Masc/Fém × 15 tranches + TTL"""
    # Ligne 0 : titre colonne 0 + tranches groupées
    row0 = t.rows[0]
    cell_text(row0.cells[0], titre_col0, bold=True, size=7, bg="C0C0C0")
    for i, (_, lbl, _) in enumerate(TRANCHES_RMA):
        cell_text(row0.cells[1 + i*2], lbl, bold=True, size=7, bg="C0C0C0")
        cell_text(row0.cells[2 + i*2], "", size=7, bg="C0C0C0")
    cell_text(row0.cells[-2], "TTL", bold=True, size=7, bg="C0C0C0")
    cell_text(row0.cells[-1], "", size=7, bg="C0C0C0")

    # Ligne 1 : M/F pour chaque tranche
    row1 = t.rows[1]
    cell_text(row1.cells[0], "", size=7, bg="C0C0C0")
    for i in range(len(TRANCHES_RMA)):
        cell_text(row1.cells[1 + i*2], "M", bold=True, size=7, bg="C0C0C0")
        cell_text(row1.cells[2 + i*2], "F", bold=True, size=7, bg="C0C0C0")
    cell_text(row1.cells[-2], "M", bold=True, size=7, bg="C0C0C0")
    cell_text(row1.cells[-1], "F", bold=True, size=7, bg="C0C0C0")


def fill_row_rma(row, libelle, serie, bg_label=None):
    """Remplit une ligne de données RMA"""
    cell_text(row.cells[0], libelle, size=7,
              align=WD_ALIGN_PARAGRAPH.LEFT, bg=bg_label)
    codes = [c for c, _, _ in TRANCHES_RMA]
    for i, code in enumerate(codes):
        cell_text(row.cells[1 + i*2], val(serie, f"m_{code}"), size=7)
        cell_text(row.cells[2 + i*2], val(serie, f"f_{code}"), size=7)
    cell_text(row.cells[-2], val(serie, "m_total"), size=7)
    cell_text(row.cells[-1], val(serie, "f_total"), size=7)


def new_table_rma(doc, n_data_rows, titre_col0=""):
    """Crée un tableau RMA standard"""
    n_cols = 1 + len(TRANCHES_RMA) * 2 + 2
    t = doc.add_table(rows=2 + n_data_rows, cols=n_cols)
    t.style = "Table Grid"
    build_header_rma(t, titre_col0)
    return t


# =============================================================
# CONSTRUCTION DES TABLEAUX
# =============================================================

def build_screening(doc, data):
    add_titre(doc, "Screening VIH — Consultants", 2)
    t = new_table_rma(doc, 5, "Indicateurs")

    lignes = [
        ("1- Nombre de consultants",                       data["tous"]),
        ("2- Screenés pour dépistage VIH",                 data["screens"]),
        ("3- Éligibles après screening",                   data["screens"]),
        ("4- Éligibles dépistés ayant reçu le résultat",   data["screens"]),
        ("5- Dépistés positifs ayant reçu le résultat",    data["positifs"]),
    ]
    for i, (lbl, serie) in enumerate(lignes):
        fill_row_rma(t.rows[2 + i], lbl, serie)

    doc.add_paragraph()


def build_conseil_depistage(doc, data):
    add_titre(doc, "I- Conseil et Dépistage", 2)
    t = new_table_rma(doc, 3, "")

    fill_row_rma(t.rows[2], "Clients dépistés ayant retiré un résultat", data["tous"])
    fill_row_rma(t.rows[3], "Résultat positif (y compris contrôles positifs)", data["positifs"])
    fill_row_rma(t.rows[4], "Résultat négatif", data["negatifs"])

    doc.add_paragraph()


def build_file_active_dotation(doc, data):
    add_titre(doc, "File Active — Dotation ARV (312)", 2)
    t = new_table_rma(doc, 4, "Indicateurs")

    fill_row_rma(t.rows[2], "312. Total File Active", data["total"])
    fill_row_rma(t.rows[3], "312.a. Dotation < 3 mois (< 90j)", data["dot_lt90"])
    fill_row_rma(t.rows[4], "312.b. Dotation 3-5 mois (90-179j)", data["dot_90_179"])
    fill_row_rma(t.rows[5], "312.c. Dotation ≥ 6 mois (≥ 180j)", data["dot_gte180"])

    doc.add_paragraph()


def build_stables(doc, data):
    add_titre(doc, "Patients Stables / Non Stables / NE (318-324)", 2)
    t = new_table_rma(doc, 6, "Indicateurs")

    fill_row_rma(t.rows[2], "318. Total Stables", data["total"])
    fill_row_rma(t.rows[3], "318.a. TARV tous 3 mois (≤ 90j)", data["a_3mois"])
    fill_row_rma(t.rows[4], "318.b. Visite clinique 6 mois (≥ 180j)", data["b_6mois"])
    fill_row_rma(t.rows[5], "318.f. Autres modèles (91-179j)", data["f_autres"])
    fill_row_rma(t.rows[6], "321. Non Stables", data["non_stables"])
    fill_row_rma(t.rows[7], "324. Non Évalués (NE)", data["non_evalues"])

    doc.add_paragraph()


def build_regimes(doc, data):
    add_titre(doc, "4. Régimes des patients sous ARV", 2)

    def get_vals(df, regime):
        row = df[df["regime"] == regime]
        if row.empty:
            return 0, 0, 0, 0
        r = row.iloc[0]
        return (int(r.get("adulte_h", 0) or 0),
                int(r.get("adulte_f", 0) or 0),
                int(r.get("enfant_h", 0) or 0),
                int(r.get("enfant_f", 0) or 0))

    # 1ère ligne adulte + enfant
    for titre, regimes in [
        ("TTT 1ère ligne Adulte (≥ 15 ans)", REGIMES_1ERE_LIGNE),
        ("TTT 1ère ligne Enfant (< 15 ans)", REGIMES_1ERE_LIGNE),
    ]:
        p = doc.add_paragraph()
        p.add_run(titre).bold = True

        t = doc.add_table(rows=len(regimes) + 2, cols=5)
        t.style = "Table Grid"

        # En-tête
        hdrs = ["Régime", "Nouveaux H", "Nouveaux F", "File active H", "File active F"]
        for i, h in enumerate(hdrs):
            cell_text(t.rows[0].cells[i], h, bold=True, bg="C0C0C0")

        for r_idx, regime in enumerate(regimes):
            nh, nf, _, _ = get_vals(data["nouveaux"], regime)
            fh, ff, _, _ = get_vals(data["file_active"], regime)

            if "Enfant" in titre:
                _, _, nh, nf = get_vals(data["nouveaux"], regime)
                _, _, fh, ff = get_vals(data["file_active"], regime)

            row = t.rows[1 + r_idx]
            cell_text(row.cells[0], regime, align=WD_ALIGN_PARAGRAPH.LEFT)
            cell_text(row.cells[1], nh)
            cell_text(row.cells[2], nf)
            cell_text(row.cells[3], fh)
            cell_text(row.cells[4], ff)

        # Total
        row_total = t.rows[-1]
        cell_text(row_total.cells[0], "TOTAL", bold=True, bg="D9D9D9")
        for ci in range(1, 5):
            total = sum(
                int(t.rows[1 + ri].cells[ci].text or 0)
                for ri in range(len(regimes))
                if t.rows[1 + ri].cells[ci].text.isdigit()
            )
            cell_text(row_total.cells[ci], total, bold=True, bg="D9D9D9")

        doc.add_paragraph()


def build_charge_virale(doc, data):
    add_titre(doc, "11a. Réalisation de charge virale (12 derniers mois)", 2)
    t = new_table_rma(doc, 4, "Indicateurs")

    fill_row_rma(t.rows[2], "1006. CV réalisée (résultats reçus — 12 mois)", data["cv_1006"])
    fill_row_rma(t.rows[3], "1009. CV ≤ 1000 copies/ml (12 mois)", data["cv_1009"])
    fill_row_rma(t.rows[4], "1012. CV > 1000 copies/ml (12 mois)", data["cv_1012"])
    fill_row_rma(t.rows[5], "1022. CV prélevée ce mois (Routine)", data["cv_1022"])

    doc.add_paragraph()


# =============================================================
# GÉNÉRATION RAPPORT
# =============================================================
def generer_rapport_rma(mois_str, output_dir=None):
    debut, fin, label = get_periode(mois_str)
    log.info(f"📄 Rapport RMA — {label} ({debut} → {fin})")

    conn = get_conn()

    log.info("  Requêtes DuckDB...")
    d_screen  = q_screening(conn, debut, fin)
    d_conseil = q_conseil_depistage(conn, debut, fin)
    d_dotation = q_file_active_dotation(conn, fin)
    d_stables  = q_stables(conn, fin)
    d_regimes  = q_regimes(conn, debut, fin)
    d_cv       = q_charge_virale(conn, debut, fin)
    conn.close()

    log.info("  Construction du document Word...")
    doc = Document()

    # Marges
    for section in doc.sections:
        section.top_margin    = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin   = Cm(1.0)
        section.right_margin  = Cm(1.0)
        section.page_width    = Cm(42)  # Paysage A3

    # Couverture
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        "RAPPORT MENSUEL DES ACTIVITÉS DE SUIVI DES\n"
        "PERSONNES INFECTÉES ET AFFECTÉES PAR LE VIH"
    )
    run.bold = True
    run.font.size = Pt(14)

    doc.add_paragraph()

    infos = [
        ("Établissement", CONFIG_ETAB["nom"]),
        ("Période",       label),
        ("District",      CONFIG_ETAB["district"]),
        ("Région",        CONFIG_ETAB["region"]),
    ]
    for lbl, v in infos:
        p = doc.add_paragraph()
        r = p.add_run(f"• {lbl} : ")
        r.bold = True
        r.font.size = Pt(11)
        p.add_run(v).font.size = Pt(11)

    doc.add_page_break()

    # Tableaux
    build_screening(doc, d_screen)
    build_conseil_depistage(doc, d_conseil)
    doc.add_page_break()
    build_file_active_dotation(doc, d_dotation)
    build_stables(doc, d_stables)
    doc.add_page_break()
    build_regimes(doc, d_regimes)
    doc.add_page_break()
    build_charge_virale(doc, d_cv)

    # Sauvegarde
    output_dir = Path(output_dir) if output_dir else Path(".")
    output_dir.mkdir(parents=True, exist_ok=True)
    mois_clean = mois_str.replace("-", "_")
    nom = f"Rapport_RMA_KoKhoua_{mois_clean}.docx"
    chemin = output_dir / nom
    doc.save(str(chemin))
    log.info(f"  ✅ Rapport RMA généré → {chemin}")
    return chemin


# =============================================================
# POINT D'ENTRÉE
# =============================================================
if __name__ == "__main__":
    parser = argparse.ArgumentParser()
    parser.add_argument(
        "--mois", type=str,
        default=(date.today() - relativedelta(months=1)).strftime("%Y-%m"),
    )
    parser.add_argument("--output", type=str, default="./rapports")
    args = parser.parse_args()
    generer_rapport_rma(args.mois, args.output)