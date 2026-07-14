# =============================================================
# KoSanté BI — rapport_sig.py
# =============================================================
# Génération automatique du Rapport Mensuel SIG DIIS
# depuis les données DuckDB + fichier Excel prélèvements
#
# USAGE :
#   python rapport_sig.py --mois 2026-06
#   python rapport_sig.py --mois 2026-06 --prelevements fichier.xlsx
# =============================================================

import duckdb
import pandas as pd
from pathlib import Path
from datetime import date, datetime
from dateutil.relativedelta import relativedelta
import sys
import argparse
import logging

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(message)s",
    datefmt="%H:%M:%S"
)
log = logging.getLogger(__name__)

# =============================================================
# CONFIG ÉTABLISSEMENT
# =============================================================
CONFIG_ETAB = {
    "nom":              "ONG KO'KHOUA",
    "num_immat":        "00132/01/",
    "district":         "Treichville-Marcory",
    "district_code":    "01",
    "region":           "Abidjan 2",
    "region_code":      "02",
    "population":       "",
    "realise_par":      "KONATE SOUMAHILA",
}

TRANCHES_AGE = ["0-4", "5-9", "10-14", "15-19", "20-24", "25-49", "50+"]

# =============================================================
# CONNEXION DUCKDB
# =============================================================
def get_conn():
    db_path = Path(__file__).parent / "data" / "kosante.duckdb"
    return duckdb.connect(str(db_path), read_only=True)


# =============================================================
# CALCUL DE LA PÉRIODE
# =============================================================
def get_periode(mois_str: str):
    """
    Retourne (debut, fin, label) pour un mois donné
    Format mois_str : 'YYYY-MM' ex: '2026-06'
    """
    d = datetime.strptime(mois_str, "%Y-%m")
    debut = date(d.year, d.month, 1)
    fin = debut + relativedelta(months=1) - relativedelta(days=1)
    label = d.strftime("%B %Y").capitalize()
    return debut, fin, label


# =============================================================
# HELPERS SQL — TRANCHES D'ÂGE
# =============================================================
def sql_tranches_age_sexe(champ_age, champ_sexe, prefix=""):
    """
    Génère les colonnes SQL pour ventilation par tranche d'âge et sexe
    """
    tranches = [
        ("t04",   f"{champ_age} BETWEEN 0 AND 4"),
        ("t59",   f"{champ_age} BETWEEN 5 AND 9"),
        ("t1014", f"{champ_age} BETWEEN 10 AND 14"),
        ("t1519", f"{champ_age} BETWEEN 15 AND 19"),
        ("t2024", f"{champ_age} BETWEEN 20 AND 24"),
        ("t2549", f"{champ_age} BETWEEN 25 AND 49"),
        ("t50p",  f"{champ_age} >= 50"),
    ]
    cols = []
    p = f"{prefix}_" if prefix else ""
    for code, cond in tranches:
        cols.append(f"COUNT(*) FILTER (WHERE ({cond}) AND {champ_sexe} = 1) AS {p}m_{code}")
        cols.append(f"COUNT(*) FILTER (WHERE ({cond}) AND {champ_sexe} = 2) AS {p}f_{code}")
    cols.append(f"COUNT(*) FILTER (WHERE {champ_sexe} = 1) AS {p}m_total")
    cols.append(f"COUNT(*) FILTER (WHERE {champ_sexe} = 2) AS {p}f_total")
    cols.append(f"COUNT(*) AS {p}total")
    return cols


def sql_tranches_age_seul(champ_age, prefix=""):
    """Tranches d'âge sans ventilation sexe"""
    tranches = [
        ("t04",   f"{champ_age} BETWEEN 0 AND 4"),
        ("t59",   f"{champ_age} BETWEEN 5 AND 9"),
        ("t1014", f"{champ_age} BETWEEN 10 AND 14"),
        ("t1519", f"{champ_age} BETWEEN 15 AND 19"),
        ("t2024", f"{champ_age} BETWEEN 20 AND 24"),
        ("t2549", f"{champ_age} BETWEEN 25 AND 49"),
        ("t50p",  f"{champ_age} >= 50"),
    ]
    cols = []
    p = f"{prefix}_" if prefix else ""
    for code, cond in tranches:
        cols.append(f"COUNT(*) FILTER (WHERE {cond}) AS {p}{code}")
    cols.append(f"COUNT(*) AS {p}total")
    return cols


# =============================================================
# NORMALISATION NUMÉRO NATIONAL
# =============================================================
NORMALISE_NUM = """
REGEXP_REPLACE(TRIM(COALESCE(CAST({col} AS VARCHAR), '')), '^0+', '')
"""

def join_regime_patient():
    """Jointure TblRegime × TblDossPatient avec normalisation NumNational"""
    return f"""
        TblRegime r
        INNER JOIN TblDossPatient p
            ON {NORMALISE_NUM.format(col='r.NumPatient')}
             = {NORMALISE_NUM.format(col='p.NumNational')}
    """


# =============================================================
# REQUÊTES DE DONNÉES
# =============================================================

def q_tableau1(conn, debut, fin):
    """I.1.1 — Soins curatifs : consultants et consultations"""

    # Consultants = dépistages CDV (âge calculé depuis DateNaiss)
    consultants = conn.execute(f"""
        SELECT
            {', '.join(sql_tranches_age_seul(
                "DATEDIFF('year', TRY_CAST(DateNaiss AS DATE), TRY_CAST(DateVisite AS DATE))"
            ))}
        FROM TblRegistreCDV
        WHERE TRY_CAST(DateVisite AS DATE) BETWEEN DATE '{debut}' AND DATE '{fin}'
          AND DateNaiss IS NOT NULL
    """).df().iloc[0]

    # Consultations = dispensations ARV
    consultations = conn.execute(f"""
        SELECT
            {', '.join(sql_tranches_age_seul("p.Age"))}
        FROM {join_regime_patient()}
        WHERE TRY_CAST(r.DateRegime AS DATE) BETWEEN DATE '{debut}' AND DATE '{fin}'
          AND r.REGIME IS NOT NULL
    """).df().iloc[0]

    return {"consultants": consultants, "consultations": consultations}


def q_tableau13a(conn, debut, fin):
    """13a — Évaluation nutritionnelle population générale"""
    return conn.execute(f"""
        SELECT
            {', '.join(sql_tranches_age_sexe("p.Age", "p.Sexe"))}
        FROM {join_regime_patient()}
        WHERE TRY_CAST(r.DateRegime AS DATE) BETWEEN DATE '{debut}' AND DATE '{fin}'
          AND r.REGIME IS NOT NULL
    """).df().iloc[0]


def q_tableau14(conn, debut, fin):
    """14 — Conseils et dépistage VIH"""
    age_expr = "DATEDIFF('year', TRY_CAST(DateNaiss AS DATE), TRY_CAST(DateVisite AS DATE))"

    # Tous les dépistages
    tous = conn.execute(f"""
        SELECT {', '.join(sql_tranches_age_sexe(age_expr, "Sexe"))}
        FROM TblRegistreCDV
        WHERE TRY_CAST(DateVisite AS DATE) BETWEEN DATE '{debut}' AND DATE '{fin}'
          AND DateNaiss IS NOT NULL
    """).df().iloc[0]

    # Positifs seulement
    positifs = conn.execute(f"""
        SELECT {', '.join(sql_tranches_age_sexe(age_expr, "Sexe"))}
        FROM TblRegistreCDV
        WHERE TRY_CAST(DateVisite AS DATE) BETWEEN DATE '{debut}' AND DATE '{fin}'
          AND DateNaiss IS NOT NULL
          AND Resultat_simple = 'Positif'
    """).df().iloc[0]

    return {"tous": tous, "positifs": positifs}


def q_tableau20(conn, debut, fin):
    """20 — Soins PVVIH"""

    # Col 3 : nouveaux patients enrôlés ce mois
    nouveaux = conn.execute(f"""
        SELECT {', '.join(sql_tranches_age_sexe("Age", "Sexe"))}
        FROM TblDossPatient
        WHERE TRY_CAST(DateAdmi AS DATE) BETWEEN DATE '{debut}' AND DATE '{fin}'
          AND TypePatient = 'Nouveau Ko''Khoua'
    """).df().iloc[0]

    # Col 4 : file active FIGÉE au dernier jour du mois de rapport
    # DatePDV >= fin_du_mois signifie le patient était encore actif
    # au 30 juin (ou dernier jour du mois)
    file_active = conn.execute(f"""
        SELECT {', '.join(sql_tranches_age_sexe("p.Age", "p.Sexe"))}
        FROM TblRegime r
        INNER JOIN TblDossPatient p
            ON REGEXP_REPLACE(TRIM(CAST(r.NumPatient AS VARCHAR)), '^0+', '')
             = REGEXP_REPLACE(TRIM(CAST(p.NumNational AS VARCHAR)), '^0+', '')
        WHERE r.Derniere_Dispensation = TRUE
          AND r.REGIME IS NOT NULL
          AND TRY_CAST(r.DatePDV AS DATE) >= DATE '{fin}'
          AND (p.DecesDate IS NULL OR p.DecesDate = '')
          AND (p.TransfDate IS NULL OR p.TransfDate = '')
          AND p.NumInc > 0
    """).df().iloc[0]

    return {"nouveaux": nouveaux, "file_active": file_active}


def q_tableau21a(conn, debut, fin, df_prelevements=None):
    """21a — Suivi biologique"""

    # Col 1 : nouveaux sous ARV ce mois
    nouveaux_arv = conn.execute(f"""
        SELECT {', '.join(sql_tranches_age_sexe("p.Age", "p.Sexe"))}
        FROM TblMiseEnRoute m
        INNER JOIN TblDossPatient p ON m.Patient = p.NumInc
        WHERE TRY_CAST(m.DateMiseTARV AS DATE) BETWEEN DATE '{debut}' AND DATE '{fin}'
    """).df().iloc[0]

    # Col 2&3 : prélèvements depuis fichier Excel
    prelevements = None
    fe_col2 = 0
    fa_col2 = 0
    if df_prelevements is not None and not df_prelevements.empty:
        df_p = df_prelevements.copy()
        df_p['Date de Prelevment'] = pd.to_datetime(
            df_p['Date de Prelevment'], errors='coerce', dayfirst=True
        )
        mask = (
            (df_p['Date de Prelevment'] >= pd.Timestamp(debut)) &
            (df_p['Date de Prelevment'] <= pd.Timestamp(fin))
        )
        df_mois = df_p[mask].copy()
        fe_col2 = int((df_mois['FE'].str.upper() == 'OUI').sum()) if 'FE' in df_mois else 0
        fa_col2 = int((df_mois['FA'].str.upper() == 'OUI').sum()) if 'FA' in df_mois else 0

        # Ventilation par sexe et tranche
        def ventiler(df, sexe_val, age_min, age_max=None):
            s = df[df['Sexe'] == sexe_val]
            if age_max:
                return len(s[(s['Age'] >= age_min) & (s['Age'] <= age_max)])
            return len(s[s['Age'] >= age_min])

        prelevements = {
            "m_t04":   ventiler(df_mois, 'M', 0, 4),
            "f_t04":   ventiler(df_mois, 'F', 0, 4),
            "m_t59":   ventiler(df_mois, 'M', 5, 9),
            "f_t59":   ventiler(df_mois, 'F', 5, 9),
            "m_t1014": ventiler(df_mois, 'M', 10, 14),
            "f_t1014": ventiler(df_mois, 'F', 10, 14),
            "m_t1519": ventiler(df_mois, 'M', 15, 19),
            "f_t1519": ventiler(df_mois, 'F', 15, 19),
            "m_t2024": ventiler(df_mois, 'M', 20, 24),
            "f_t2024": ventiler(df_mois, 'F', 20, 24),
            "m_t2549": ventiler(df_mois, 'M', 25, 49),
            "f_t2549": ventiler(df_mois, 'F', 25, 49),
            "m_50p":   ventiler(df_mois, 'M', 50),
            "f_50p":   ventiler(df_mois, 'F', 50),
            "m_total": len(df_mois[df_mois['Sexe'] == 'M']),
            "f_total": len(df_mois[df_mois['Sexe'] == 'F']),
            "total":   len(df_mois),
            "fe":      fe_col2,
            "fa":      fa_col2,
        }

    # Col 4 : CV supprimée ce mois
    cv_supprimee = conn.execute(f"""
        SELECT {', '.join(sql_tranches_age_sexe("p.Age", "p.Sexe"))}
        FROM TblChargesVirales cv
        INNER JOIN TblDossPatient p ON cv.Patient = p.NumInc
        WHERE TRY_CAST(cv.DatePrelev AS DATE) BETWEEN DATE '{debut}' AND DATE '{fin}'
          AND cv.CVcopies <= 1000
          AND cv.CVcopies IS NOT NULL
    """).df().iloc[0]

    return {
        "nouveaux_arv": nouveaux_arv,
        "prelevements": prelevements,
        "cv_supprimee": cv_supprimee,
    }


def q_tableau21b(conn):
    """21b — Patients stables"""
    return conn.execute("""
        SELECT
            COUNT(*) FILTER (
                WHERE sc.StatutCV = 'Stable' AND p.Age < 15
            ) AS enfant,
            COUNT(*) FILTER (
                WHERE sc.StatutCV = 'Stable' AND p.Age >= 15
            ) AS adulte,
            COUNT(*) FILTER (WHERE sc.StatutCV = 'Stable') AS total
        FROM StatutCV_Patient sc
        INNER JOIN file_active fa ON sc.NumInc = fa.NumInc
        INNER JOIN TblDossPatient p ON fa.NumInc = p.NumInc
        WHERE fa.StatutFile = 'Actif'
    """).df().iloc[0]


def q_tableau22(conn, debut, fin):
    """22 — Évaluation nutritionnelle PVVIH"""
    return conn.execute(f"""
        SELECT {', '.join(sql_tranches_age_sexe("p.Age", "p.Sexe"))}
        FROM {join_regime_patient()}
        INNER JOIN StatutCV_Patient sc ON p.NumInc = sc.NumInc
        WHERE TRY_CAST(r.DateRegime AS DATE) BETWEEN DATE '{debut}' AND DATE '{fin}'
          AND r.REGIME IS NOT NULL
    """).df().iloc[0]


def q_tableau23(conn, debut, fin):
    """23 — Traitement ARV"""

    # Col 1 : nouveaux dépistés enrôlés
    nouveaux = conn.execute(f"""
        SELECT {', '.join(sql_tranches_age_sexe("Age", "Sexe"))}
        FROM TblDossPatient
        WHERE TRY_CAST(DateAdmi AS DATE) BETWEEN DATE '{debut}' AND DATE '{fin}'
          AND TypePatient = 'Nouveau Ko''Khoua'
    """).df().iloc[0]

    # Col 4 : initiés ARV ce mois
    inities = conn.execute(f"""
        SELECT {', '.join(sql_tranches_age_sexe("p.Age", "p.Sexe"))}
        FROM TblMiseEnRoute m
        INNER JOIN TblDossPatient p ON m.Patient = p.NumInc
        WHERE TRY_CAST(m.DateMiseTARV AS DATE) BETWEEN DATE '{debut}' AND DATE '{fin}'
    """).df().iloc[0]

    # Col 5 : file active FIGÉE au dernier jour du mois de rapport
    file_active = conn.execute(f"""
        SELECT {', '.join(sql_tranches_age_sexe("p.Age", "p.Sexe"))}
        FROM TblRegime r
        INNER JOIN TblDossPatient p
            ON REGEXP_REPLACE(TRIM(CAST(r.NumPatient AS VARCHAR)), '^0+', '')
             = REGEXP_REPLACE(TRIM(CAST(p.NumNational AS VARCHAR)), '^0+', '')
        WHERE r.Derniere_Dispensation = TRUE
          AND r.REGIME IS NOT NULL
          AND TRY_CAST(r.DatePDV AS DATE) >= DATE '{fin}'
          AND (p.DecesDate IS NULL OR p.DecesDate = '')
          AND (p.TransfDate IS NULL OR p.TransfDate = '')
          AND p.NumInc > 0
    """).df().iloc[0]

    return {"nouveaux": nouveaux, "inities": inities, "file_active": file_active}


def q_tableau24(conn, debut, fin):
    """24 — TB/VIH"""
    return conn.execute(f"""
        SELECT {', '.join(sql_tranches_age_sexe("p.Age", "p.Sexe"))}
        FROM {join_regime_patient()}
        WHERE TRY_CAST(r.DateRegime AS DATE) BETWEEN DATE '{debut}' AND DATE '{fin}'
          AND r.REGIME IS NOT NULL
    """).df().iloc[0]


def q_tableau26(conn, debut, fin):
    """26 — Absents RDV, PDV, Décès"""

    # Col 1 : RDV manqués ce mois
    rdv = conn.execute(f"""
        SELECT {', '.join(sql_tranches_age_sexe("p.Age", "p.Sexe"))}
        FROM patients_rdv_manque rm
        INNER JOIN TblDossPatient p ON rm.NumInc = p.NumInc
        WHERE TRY_CAST(rm.DateProchainRdv AS DATE) BETWEEN DATE '{debut}' AND DATE '{fin}'
    """).df().iloc[0]

    # Col 2 : PDV ce mois
    pdv = conn.execute(f"""
        SELECT {', '.join(sql_tranches_age_sexe("p.Age", "p.Sexe"))}
        FROM attrition a
        INNER JOIN TblDossPatient p ON a.NumInc = p.NumInc
        WHERE a.categorie_attrition = 'Perdu de vue'
          AND TRY_CAST(a.DatePDV AS DATE) BETWEEN DATE '{debut}' AND DATE '{fin}'
    """).df().iloc[0]

    # Col 3 : Décès ce mois
    deces = conn.execute(f"""
        SELECT {', '.join(sql_tranches_age_sexe("p.Age", "p.Sexe"))}
        FROM attrition a
        INNER JOIN TblDossPatient p ON a.NumInc = p.NumInc
        WHERE a.categorie_attrition = 'Décès'
          AND TRY_CAST(a.DecesDate AS DATE) BETWEEN DATE '{debut}' AND DATE '{fin}'
    """).df().iloc[0]

    return {"rdv": rdv, "pdv": pdv, "deces": deces}


def q_tableau27(conn, debut, fin):
    """27 — Cohorte 12 mois"""
    debut_cohorte = debut - relativedelta(months=12)
    fin_cohorte   = fin   - relativedelta(months=12)

    # Total cohorte
    cohorte = conn.execute(f"""
        SELECT
            COUNT(*) FILTER (WHERE p.Age < 15)  AS enfant,
            COUNT(*) FILTER (WHERE p.Age >= 15) AS adulte,
            COUNT(*) AS total
        FROM TblMiseEnRoute m
        INNER JOIN TblDossPatient p ON m.Patient = p.NumInc
        WHERE TRY_CAST(m.DateMiseTARV AS DATE)
              BETWEEN DATE '{debut_cohorte}' AND DATE '{fin_cohorte}'
    """).df().iloc[0]

    # Encore actifs
    actifs = conn.execute(f"""
        SELECT
            COUNT(*) FILTER (WHERE p.Age < 15)  AS enfant,
            COUNT(*) FILTER (WHERE p.Age >= 15) AS adulte,
            COUNT(*) AS total
        FROM TblMiseEnRoute m
        INNER JOIN TblDossPatient p ON m.Patient = p.NumInc
        INNER JOIN file_active fa ON m.Patient = fa.NumInc
        WHERE TRY_CAST(m.DateMiseTARV AS DATE)
              BETWEEN DATE '{debut_cohorte}' AND DATE '{fin_cohorte}'
          AND fa.StatutFile = 'Actif'
    """).df().iloc[0]

    # CV prélevée dans fenêtre ±3 mois autour des 12 mois
    cv_prelev = conn.execute(f"""
        SELECT
            COUNT(DISTINCT m.Patient) FILTER (
                WHERE p.Age < 15
            ) AS enfant,
            COUNT(DISTINCT m.Patient) FILTER (
                WHERE p.Age >= 15
            ) AS adulte,
            COUNT(DISTINCT m.Patient) AS total
        FROM TblMiseEnRoute m
        INNER JOIN TblDossPatient p ON m.Patient = p.NumInc
        INNER JOIN TblChargesVirales cv ON m.Patient = cv.Patient
        WHERE TRY_CAST(m.DateMiseTARV AS DATE)
              BETWEEN DATE '{debut_cohorte}' AND DATE '{fin_cohorte}'
          AND TRY_CAST(cv.DatePrelev AS DATE)
              BETWEEN TRY_CAST(m.DateMiseTARV AS DATE) + INTERVAL '9 months'
                  AND TRY_CAST(m.DateMiseTARV AS DATE) + INTERVAL '15 months'
    """).df().iloc[0]

    # CV supprimée (< 1000 copies)
    cv_supp = conn.execute(f"""
        SELECT
            COUNT(DISTINCT m.Patient) FILTER (
                WHERE p.Age < 15
            ) AS enfant,
            COUNT(DISTINCT m.Patient) FILTER (
                WHERE p.Age >= 15
            ) AS adulte,
            COUNT(DISTINCT m.Patient) AS total
        FROM TblMiseEnRoute m
        INNER JOIN TblDossPatient p ON m.Patient = p.NumInc
        INNER JOIN TblChargesVirales cv ON m.Patient = cv.Patient
        WHERE TRY_CAST(m.DateMiseTARV AS DATE)
              BETWEEN DATE '{debut_cohorte}' AND DATE '{fin_cohorte}'
          AND TRY_CAST(cv.DatePrelev AS DATE)
              BETWEEN TRY_CAST(m.DateMiseTARV AS DATE) + INTERVAL '9 months'
                  AND TRY_CAST(m.DateMiseTARV AS DATE) + INTERVAL '15 months'
          AND cv.CVcopies <= 1000
    """).df().iloc[0]

    return {
        "cohorte": cohorte,
        "actifs": actifs,
        "cv_prelev": cv_prelev,
        "cv_supp": cv_supp,
    }


# =============================================================
# HELPERS DOCX
# =============================================================
def set_cell_bg(cell, hex_color):
    """Couleur de fond d'une cellule"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def cell_text(cell, text, bold=False, size=9, align=WD_ALIGN_PARAGRAPH.CENTER, color=None):
    """Écrire du texte dans une cellule"""
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(str(text) if text is not None else "")
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def val(row, col):
    """Retourne une valeur ou 0 si None/NaN"""
    try:
        v = row[col]
        if pd.isna(v):
            return 0
        return int(v)
    except (KeyError, TypeError):
        return 0


def add_section_title(doc, text, level=1):
    """Ajoute un titre de section"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12 if level == 1 else 10)
    if level == 1:
        p.paragraph_format.space_before = Pt(12)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), '000000')
        shading_elm.set(qn('w:val'), 'clear')
        p._p.get_or_add_pPr().append(shading_elm)
    return p


def add_tableau_title(doc, text):
    """Titre d'un tableau"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(10)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)


# =============================================================
# CONSTRUCTION DES TABLEAUX DOCX
# =============================================================

def tableau1_soins_curatifs(doc, data):
    """I.1.1 — Activités de soins curatifs"""
    add_tableau_title(doc, "Tableau 1 : Activités de soins curatifs")

    headers_age = ["0-4 ans", "5-09 ans", "10-14 ans", "15-19 ans",
                   "20-24 ans", "25-49 ans", "50 ans et plus", "Total"]
    cols_age = ["t04", "t59", "t1014", "t1519", "t2024", "t2549", "t50p", "total"]

    t = doc.add_table(rows=3, cols=9)
    t.style = 'Table Grid'

    # En-tête
    hdr = t.rows[0]
    cell_text(hdr.cells[0], "Activités", bold=True, size=8)
    set_cell_bg(hdr.cells[0], 'C0C0C0')
    for i, h in enumerate(headers_age):
        cell_text(hdr.cells[i+1], h, bold=True, size=8)
        set_cell_bg(hdr.cells[i+1], 'C0C0C0')

    # Ligne consultants
    row = t.rows[1]
    cell_text(row.cells[0], "Nombre de consultants", size=8,
              align=WD_ALIGN_PARAGRAPH.LEFT)
    for i, c in enumerate(cols_age):
        cell_text(row.cells[i+1], val(data["consultants"], c), size=8)

    # Ligne consultations
    row = t.rows[2]
    cell_text(row.cells[0], "Nombre de consultations", size=8,
              align=WD_ALIGN_PARAGRAPH.LEFT)
    for i, c in enumerate(cols_age):
        cell_text(row.cells[i+1], val(data["consultations"], c), size=8)

    doc.add_paragraph()


def tableau14_depistage(doc, data):
    """Tableau 14 — Conseils et dépistage VIH"""
    add_tableau_title(doc, "Tableau 14 : Conseils et dépistage (hormis PTME)")

    sexes = ["M", "F"]
    tranches = ["t04", "t59", "t1014", "t1519", "t2024", "t2549", "t50p"]
    labels = ["0-4", "5-09", "10-14", "15-19", "20-24", "25-49", "50+"]

    # Nombre de colonnes : 1 (libellé) + 7*2 (tranches M/F) + 2 (Total M/F)
    ncols = 1 + 7*2 + 2
    t = doc.add_table(rows=5, cols=ncols)
    t.style = 'Table Grid'

    # Ligne 0 : libellé activité | tranches d'âge groupées | Total
    row0 = t.rows[0]
    cell_text(row0.cells[0], "Activités", bold=True, size=7)
    set_cell_bg(row0.cells[0], 'C0C0C0')
    for i, lbl in enumerate(labels):
        col_idx = 1 + i*2
        cell_text(row0.cells[col_idx], lbl, bold=True, size=7)
        set_cell_bg(row0.cells[col_idx], 'C0C0C0')
        cell_text(row0.cells[col_idx+1], "", size=7)
        set_cell_bg(row0.cells[col_idx+1], 'C0C0C0')
    cell_text(row0.cells[-2], "TOTAL", bold=True, size=7)
    set_cell_bg(row0.cells[-2], 'C0C0C0')
    cell_text(row0.cells[-1], "", size=7)
    set_cell_bg(row0.cells[-1], 'C0C0C0')

    # Ligne 1 : M/F pour chaque tranche
    row1 = t.rows[1]
    cell_text(row1.cells[0], "", size=7)
    set_cell_bg(row1.cells[0], 'C0C0C0')
    for i in range(7):
        cell_text(row1.cells[1+i*2], "M", bold=True, size=7)
        set_cell_bg(row1.cells[1+i*2], 'C0C0C0')
        cell_text(row1.cells[2+i*2], "F", bold=True, size=7)
        set_cell_bg(row1.cells[2+i*2], 'C0C0C0')
    cell_text(row1.cells[-2], "M", bold=True, size=7)
    set_cell_bg(row1.cells[-2], 'C0C0C0')
    cell_text(row1.cells[-1], "F", bold=True, size=7)
    set_cell_bg(row1.cells[-1], 'C0C0C0')

    def fill_row(row_idx, libelle, serie):
        row = t.rows[row_idx]
        cell_text(row.cells[0], libelle, size=7, align=WD_ALIGN_PARAGRAPH.LEFT)
        for i, tr in enumerate(tranches):
            cell_text(row.cells[1+i*2], val(serie, f"m_{tr}"), size=7)
            cell_text(row.cells[2+i*2], val(serie, f"f_{tr}"), size=7)
        cell_text(row.cells[-2], val(serie, "m_total"), size=7)
        cell_text(row.cells[-1], val(serie, "f_total"), size=7)

    fill_row(2, "Nombre de clients conseillés", data["tous"])
    fill_row(3, "Nombre de clients conseillés et dépistés ayant reçu le résultat", data["tous"])
    fill_row(4, "Nombre de clients dépistés positif", data["positifs"])

    doc.add_paragraph()


def tableau_sexe_age_generique(doc, titre, lignes_data, labels_lignes):
    """
    Tableau générique avec ventilation sexe × tranche d'âge
    lignes_data : liste de Series pandas
    labels_lignes : liste de str
    """
    add_tableau_title(doc, titre)

    tranches = ["t04", "t59", "t1014", "t1519", "t2024", "t2549", "t50p"]
    labels_t = ["0-4", "5-09", "10-14", "15-19", "20-24", "25-49", "50+"]

    ncols = 1 + 7*2 + 2
    nrows = 2 + len(lignes_data)
    t = doc.add_table(rows=nrows, cols=ncols)
    t.style = 'Table Grid'

    # En-tête ligne 0
    row0 = t.rows[0]
    set_cell_bg(row0.cells[0], 'C0C0C0')
    for i, lbl in enumerate(labels_t):
        cell_text(row0.cells[1+i*2], lbl, bold=True, size=7)
        set_cell_bg(row0.cells[1+i*2], 'C0C0C0')
        set_cell_bg(row0.cells[2+i*2], 'C0C0C0')
    cell_text(row0.cells[-2], "TOTAL", bold=True, size=7)
    set_cell_bg(row0.cells[-2], 'C0C0C0')
    set_cell_bg(row0.cells[-1], 'C0C0C0')

    # En-tête ligne 1 (M/F)
    row1 = t.rows[1]
    set_cell_bg(row1.cells[0], 'C0C0C0')
    for i in range(7):
        cell_text(row1.cells[1+i*2], "M", bold=True, size=7)
        set_cell_bg(row1.cells[1+i*2], 'C0C0C0')
        cell_text(row1.cells[2+i*2], "F", bold=True, size=7)
        set_cell_bg(row1.cells[2+i*2], 'C0C0C0')
    cell_text(row1.cells[-2], "M", bold=True, size=7)
    set_cell_bg(row1.cells[-2], 'C0C0C0')
    cell_text(row1.cells[-1], "F", bold=True, size=7)
    set_cell_bg(row1.cells[-1], 'C0C0C0')

    # Données
    for r_idx, (serie, libelle) in enumerate(zip(lignes_data, labels_lignes)):
        row = t.rows[2 + r_idx]
        cell_text(row.cells[0], libelle, size=7, align=WD_ALIGN_PARAGRAPH.LEFT)
        for i, tr in enumerate(tranches):
            cell_text(row.cells[1+i*2], val(serie, f"m_{tr}"), size=7)
            cell_text(row.cells[2+i*2], val(serie, f"f_{tr}"), size=7)
        cell_text(row.cells[-2], val(serie, "m_total"), size=7)
        cell_text(row.cells[-1], val(serie, "f_total"), size=7)

    doc.add_paragraph()


def tableau21b_stables(doc, data):
    """21b — Patients stables"""
    add_tableau_title(doc, "Tableau 21b : Suivi des patients stables et soins différenciés")

    t = doc.add_table(rows=3, cols=3)
    t.style = 'Table Grid'

    headers = ["", "Patients VIH+ classés stables", "Patients enrôlés soins différenciés"]
    for i, h in enumerate(headers):
        cell_text(t.rows[0].cells[i], h, bold=True, size=8)
        set_cell_bg(t.rows[0].cells[i], 'C0C0C0')

    cell_text(t.rows[1].cells[0], "Enfant < 15 ans", size=8, align=WD_ALIGN_PARAGRAPH.LEFT)
    cell_text(t.rows[1].cells[1], val(data, "enfant"), size=8)
    cell_text(t.rows[1].cells[2], val(data, "enfant"), size=8)

    cell_text(t.rows[2].cells[0], "Adulte ≥ 15 ans", size=8, align=WD_ALIGN_PARAGRAPH.LEFT)
    cell_text(t.rows[2].cells[1], val(data, "adulte"), size=8)
    cell_text(t.rows[2].cells[2], val(data, "adulte"), size=8)

    doc.add_paragraph()


def tableau26_absents(doc, data):
    """26 — Patients n'ayant pas cherché leurs ARV"""
    add_tableau_title(doc, "Tableau 26 : Patients qui ne sont pas venus chercher leurs ARV")

    tranches = ["t04", "t59", "t1014", "t1519", "t2024", "t2549", "t50p"]
    labels_t = ["0-4", "5-09", "10-14", "15-19", "20-24", "25-49", "50+"]
    lignes = [
        ("rdv",   "1. Attendus RDV mais pas venus"),
        ("pdv",   "2. Perdus de vue"),
        ("deces", "3. Décédés"),
    ]

    ncols = 1 + 7*2 + 2
    t = doc.add_table(rows=2 + len(lignes), cols=ncols)
    t.style = 'Table Grid'

    row0 = t.rows[0]
    set_cell_bg(row0.cells[0], 'C0C0C0')
    for i, lbl in enumerate(labels_t):
        cell_text(row0.cells[1+i*2], lbl, bold=True, size=7)
        set_cell_bg(row0.cells[1+i*2], 'C0C0C0')
        set_cell_bg(row0.cells[2+i*2], 'C0C0C0')
    cell_text(row0.cells[-2], "TOTAL", bold=True, size=7)
    set_cell_bg(row0.cells[-2], 'C0C0C0')
    set_cell_bg(row0.cells[-1], 'C0C0C0')

    row1 = t.rows[1]
    set_cell_bg(row1.cells[0], 'C0C0C0')
    for i in range(7):
        cell_text(row1.cells[1+i*2], "M", bold=True, size=7)
        set_cell_bg(row1.cells[1+i*2], 'C0C0C0')
        cell_text(row1.cells[2+i*2], "F", bold=True, size=7)
        set_cell_bg(row1.cells[2+i*2], 'C0C0C0')
    cell_text(row1.cells[-2], "M", bold=True, size=7)
    set_cell_bg(row1.cells[-2], 'C0C0C0')
    cell_text(row1.cells[-1], "F", bold=True, size=7)
    set_cell_bg(row1.cells[-1], 'C0C0C0')

    for r_idx, (key, libelle) in enumerate(lignes):
        row = t.rows[2 + r_idx]
        serie = data[key]
        cell_text(row.cells[0], libelle, size=7, align=WD_ALIGN_PARAGRAPH.LEFT)
        for i, tr in enumerate(tranches):
            cell_text(row.cells[1+i*2], val(serie, f"m_{tr}"), size=7)
            cell_text(row.cells[2+i*2], val(serie, f"f_{tr}"), size=7)
        cell_text(row.cells[-2], val(serie, "m_total"), size=7)
        cell_text(row.cells[-1], val(serie, "f_total"), size=7)

    doc.add_paragraph()


def tableau27_cohorte(doc, data):
    """27 — Suivi de cohorte"""
    add_tableau_title(doc, "Tableau 27 : Suivi de cohorte (12 mois)")

    headers = [
        "",
        "Total cohorte 12 mois (Z)",
        "En vie et sous ARV (V)",
        "CV demandée 12 mois",
        "2ème ligne (J)",
        "CV < 1000 copies/ml",
    ]
    lignes = [
        ("Enfant", "enfant"),
        ("Adulte", "adulte"),
    ]

    t = doc.add_table(rows=3, cols=6)
    t.style = 'Table Grid'

    for i, h in enumerate(headers):
        cell_text(t.rows[0].cells[i], h, bold=True, size=8)
        set_cell_bg(t.rows[0].cells[i], 'C0C0C0')

    for r_idx, (label, key) in enumerate(lignes):
        row = t.rows[1 + r_idx]
        cell_text(row.cells[0], label, size=8, align=WD_ALIGN_PARAGRAPH.LEFT)
        cell_text(row.cells[1], val(data["cohorte"], key), size=8)
        cell_text(row.cells[2], val(data["actifs"], key), size=8)
        cell_text(row.cells[3], val(data["cv_prelev"], key), size=8)
        cell_text(row.cells[4], "", size=8)  # 2ème ligne → vide
        cell_text(row.cells[5], val(data["cv_supp"], key), size=8)

    doc.add_paragraph()


# =============================================================
# GÉNÉRATION DU DOCUMENT WORD
# =============================================================
def generer_rapport(mois_str: str, df_prelevements=None, output_dir=None):
    """Génère le rapport SIG pour le mois donné"""

    debut, fin, label = get_periode(mois_str)
    log.info(f"📄 Génération rapport — {label} ({debut} → {fin})")

    conn = get_conn()

    # Récupération des données
    log.info("  Requêtes DuckDB...")
    d1  = q_tableau1(conn, debut, fin)
    d13 = q_tableau13a(conn, debut, fin)
    d14 = q_tableau14(conn, debut, fin)
    d20 = q_tableau20(conn, debut, fin)
    d21a = q_tableau21a(conn, debut, fin, df_prelevements)
    d21b = q_tableau21b(conn)
    d22 = q_tableau22(conn, debut, fin)
    d23 = q_tableau23(conn, debut, fin)
    d24 = q_tableau24(conn, debut, fin)
    d26 = q_tableau26(conn, debut, fin)
    d27 = q_tableau27(conn, debut, fin)
    conn.close()

    log.info("  Construction du document Word...")

    doc = Document()

    # Marges
    section = doc.sections[0]
    section.top_margin    = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin   = Cm(1.5)
    section.right_margin  = Cm(1.5)

    # ── PAGE DE COUVERTURE ─────────────────────────────────
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("MINISTÈRE DE LA SANTÉ ET DE L'HYGIÈNE PUBLIQUE")
    run.bold = True
    run.font.size = Pt(11)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("DIRECTION DE L'INFORMATIQUE ET DE L'INFORMATION SANITAIRE (DIIS)")
    r2.font.size = Pt(10)

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run("SYSTÈME D'INFORMATION DE GESTION (SIG)")
    r3.bold = True
    r3.font.size = Pt(12)

    doc.add_paragraph()

    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r4 = p4.add_run("RAPPORT MENSUEL DE L'ÉTABLISSEMENT SANITAIRE PRIMAIRE")
    r4.bold = True
    r4.font.size = Pt(14)

    doc.add_paragraph()

    # Infos établissement — modifiables
    infos = [
        ("Rapport du mois de",        label),
        ("Établissement sanitaire de", CONFIG_ETAB["nom"]),
        ("Numéro d'immatriculation",   CONFIG_ETAB["num_immat"]),
        ("District Sanitaire de",
         f"{CONFIG_ETAB['district']}   Code : {CONFIG_ETAB['district_code']}"),
        ("Région Sanitaire du",
         f"{CONFIG_ETAB['region']}   Code : {CONFIG_ETAB['region_code']}"),
        ("Population desservie",       CONFIG_ETAB["population"]),
        ("Rapport réalisé par",        CONFIG_ETAB["realise_par"]),
    ]
    for label_info, valeur in infos:
        p = doc.add_paragraph()
        r = p.add_run(f"• {label_info} : ")
        r.bold = True
        r.font.size = Pt(10)
        p.add_run(valeur).font.size = Pt(10)

    doc.add_page_break()

    # ── SECTION I — ACTIVITÉS ─────────────────────────────
    add_section_title(doc, "I. ACTIVITÉS")

    p = doc.add_paragraph()
    r = p.add_run("I.1. Activités de consultations et de soins")
    r.bold = True
    r.font.size = Pt(11)

    tableau1_soins_curatifs(doc, d1)

    # ── TABLEAU 13a ────────────────────────────────────────
    p = doc.add_paragraph()
    r = p.add_run("I.2.4. Activités d'évaluation nutritionnelle")
    r.bold = True
    r.font.size = Pt(11)

    tableau_sexe_age_generique(
        doc,
        "Tableau 13a : État nutritionnel",
        [d13],
        ["Nb personnes ayant bénéficié d'une évaluation nutritionnelle"]
    )

    # ── TABLEAU 14 ─────────────────────────────────────────
    p = doc.add_paragraph()
    r = p.add_run("I.3.1. Conseil dépistage du VIH")
    r.bold = True
    r.font.size = Pt(11)

    tableau14_depistage(doc, d14)

    # ── TABLEAU 20 ─────────────────────────────────────────
    p = doc.add_paragraph()
    r = p.add_run("I.3.3. Activités de suivi des patients séropositifs au VIH")
    r.bold = True
    r.font.size = Pt(11)

    tableau_sexe_age_generique(
        doc,
        "Tableau 20 : Soins des PVVIH — Nouveaux patients et patients suivis",
        [d20["nouveaux"], d20["file_active"]],
        [
            "Dépistés positifs enrôlés dans les soins VIH ce mois",
            "Patients VIH+ ayant reçu les soins VIH (y compris ARV)",
        ]
    )

    # ── TABLEAU 21a ────────────────────────────────────────
    p = doc.add_paragraph()
    r = p.add_run("Tableau 21a : Suivi biologique des PVVIH")
    r.bold = True
    r.font.size = Pt(11)

    lignes_21a = [d21a["nouveaux_arv"], d21a["cv_supprimee"]]
    labels_21a = [
        "Bilan initial (nouveaux sous ARV ce mois)",
        "CV ≤ 1000 copies/ml ce mois",
    ]
    if d21a["prelevements"]:
        lignes_21a.insert(1, pd.Series(d21a["prelevements"]))
        lignes_21a.insert(2, pd.Series(d21a["prelevements"]))
        labels_21a.insert(1, "Bilan biologique de suivi (prélèvements du mois)")
        labels_21a.insert(2, "Résultats CV reçus ce mois")

    tableau_sexe_age_generique(doc, "Tableau 21a", lignes_21a, labels_21a)

    # ── TABLEAU 21b ────────────────────────────────────────
    tableau21b_stables(doc, d21b)

    # ── TABLEAU 22 ─────────────────────────────────────────
    tableau_sexe_age_generique(
        doc,
        "Tableau 22 : Évaluation nutritionnelle des PVVIH",
        [d22],
        ["Patients VIH+ ayant bénéficié d'une évaluation nutritionnelle"]
    )

    # ── TABLEAU 23 ─────────────────────────────────────────
    tableau_sexe_age_generique(
        doc,
        "Tableau 23 : Patients sous Traitement ARV",
        [d23["nouveaux"], d23["inities"], d23["file_active"]],
        [
            "Dépistés positifs et mis sous ARV ce mois",
            "Nouvellement initiés sous ARV ce mois",
            "Patients VIH+ sous ARV (file active)",
        ]
    )

    # ── TABLEAU 24 ─────────────────────────────────────────
    tableau_sexe_age_generique(
        doc,
        "Tableau 24 : Soins préventifs — Coïnfection TB/VIH",
        [d24],
        ["Patients VIH+ ayant bénéficié d'une recherche active TB"]
    )

    # ── TABLEAU 26 ─────────────────────────────────────────
    tableau26_absents(doc, d26)

    # ── TABLEAU 27 ─────────────────────────────────────────
    tableau27_cohorte(doc, d27)

    # ── SAUVEGARDE ─────────────────────────────────────────
    output_dir = Path(output_dir) if output_dir else Path(".")
    output_dir.mkdir(parents=True, exist_ok=True)
    mois_clean = mois_str.replace("-", "_")
    nom_fichier = f"Rapport_SIG_{CONFIG_ETAB['district'].replace(' ', '_')}_{mois_clean}.docx"
    chemin = output_dir / nom_fichier
    doc.save(str(chemin))

    log.info(f"  ✅ Rapport généré → {chemin}")
    return chemin


# =============================================================
# POINT D'ENTRÉE
# =============================================================
if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Génération Rapport SIG Ko'Khoua")
    parser.add_argument(
        "--mois", type=str,
        default=(date.today() - relativedelta(months=1)).strftime("%Y-%m"),
        help="Mois du rapport format YYYY-MM (défaut: mois précédent)"
    )
    parser.add_argument(
        "--prelevements", type=str, default=None,
        help="Chemin vers le fichier Excel des prélèvements (optionnel)"
    )
    parser.add_argument(
        "--output", type=str, default="./rapports",
        help="Dossier de sortie du rapport"
    )
    args = parser.parse_args()

    df_prev = None
    if args.prelevements:
        log.info(f"  📂 Lecture fichier prélèvements : {args.prelevements}")
        df_prev = pd.read_excel(args.prelevements, sheet_name="CNTS", header=3)

    generer_rapport(args.mois, df_prev, args.output)